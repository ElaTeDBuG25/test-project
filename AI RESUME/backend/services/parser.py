import io
import re
from typing import Optional

from PyPDF2 import PdfReader
from docx import Document


def _extract_from_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    texts = []
    for page in reader.pages:
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        if t:
            texts.append(t)
    return "\n".join(texts)


def _extract_from_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text]
    # Tables
    for table in doc.tables:
        for row in table.rows:
            parts.extend([cell.text for cell in row.cells if cell.text])
    return "\n".join(parts)


def _extract_from_text(data: bytes, encoding: Optional[str] = None) -> str:
    encodings = [encoding] if encoding else []
    encodings += ["utf-8", "latin-1", "cp1252"]
    for enc in encodings:
        try:
            return data.decode(enc)
        except Exception:
            continue
    return data.decode("utf-8", errors="ignore")


def extract_text_from_bytes(data: bytes, filename: str) -> str:
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _extract_from_pdf(data)
    if name.endswith(".docx"):
        return _extract_from_docx(data)
    if name.endswith(".txt"):
        return _extract_from_text(data)
    # Try best-effort text for unknown types
    text = _extract_from_text(data)
    # Collapse excessive whitespace
    text = re.sub(r"\s+", " ", text).strip()
    return text
